"""Tests for PDF and DOCX converters."""

import tempfile
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from tomd import tomd


def test_tomd_with_pdf_file():
    """Test converting a PDF file to Markdown."""
    # Create a simple PDF file
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)

    # Add text to the page
    from pypdf.generic import ContentStream, NameObject, TextStringObject

    with tempfile.NamedTemporaryFile(mode="wb", suffix=".pdf", delete=False) as f:
        writer.write(f)
        temp_path = f.name

    try:
        result = tomd(temp_path)
        # PDF should return some content (even if empty page)
        assert isinstance(result, str)
    finally:
        Path(temp_path).unlink()


def test_tomd_with_docx_file():
    """Test converting a DOCX file to Markdown."""
    # Create a simple DOCX file
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph with **bold** text.")

    # Add a paragraph with bold text
    p = doc.add_paragraph()
    run = p.add_run("Bold text")
    run.bold = True

    # Add a list
    doc.add_paragraph("Item 1")
    doc.add_paragraph("Item 2")

    with tempfile.NamedTemporaryFile(mode="wb", suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    try:
        result = tomd(temp_path)

        # Check that key elements are converted
        assert "# Test Document" in result
        assert "test paragraph" in result
        assert "**Bold text**" in result
        assert "Item 1" in result
        assert "Item 2" in result
    finally:
        Path(temp_path).unlink()


def test_tomd_with_docx_table():
    """Test converting a DOCX file with a table to Markdown."""
    doc = Document()
    doc.add_heading("Document with Table", level=1)

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Row 1 Col 1"
    table.rows[1].cells[1].text = "Row 1 Col 2"
    table.rows[2].cells[0].text = "Row 2 Col 1"
    table.rows[2].cells[1].text = "Row 2 Col 2"

    with tempfile.NamedTemporaryFile(mode="wb", suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    try:
        result = tomd(temp_path)

        # Check that table is converted to markdown format
        assert "Header 1" in result
        assert "Header 2" in result
        assert "Row 1 Col 1" in result
        assert "|" in result  # Markdown table separator
        assert "---" in result  # Markdown table header separator
    finally:
        Path(temp_path).unlink()


def test_tomd_with_docx_headings():
    """Test converting DOCX headings to Markdown."""
    doc = Document()
    doc.add_heading("Heading 1", level=1)
    doc.add_heading("Heading 2", level=2)
    doc.add_heading("Heading 3", level=3)

    with tempfile.NamedTemporaryFile(mode="wb", suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    try:
        result = tomd(temp_path)

        # Check heading levels
        assert "# Heading 1" in result
        assert "## Heading 2" in result
        assert "### Heading 3" in result
    finally:
        Path(temp_path).unlink()


def test_tomd_with_docx_formatting():
    """Test converting DOCX text formatting to Markdown."""
    doc = Document()

    # Bold text
    p1 = doc.add_paragraph()
    run1 = p1.add_run("Bold text")
    run1.bold = True

    # Italic text
    p2 = doc.add_paragraph()
    run2 = p2.add_run("Italic text")
    run2.italic = True

    # Bold and italic
    p3 = doc.add_paragraph()
    run3 = p3.add_run("Bold and italic")
    run3.bold = True
    run3.italic = True

    with tempfile.NamedTemporaryFile(mode="wb", suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    try:
        result = tomd(temp_path)

        # Check formatting
        assert "**Bold text**" in result
        assert "*Italic text*" in result
        assert "***Bold and italic***" in result
    finally:
        Path(temp_path).unlink()


def test_pdf_extension_supported():
    """Test that .pdf extension is recognized."""
    with tempfile.NamedTemporaryFile(mode="wb", suffix=".pdf", delete=False) as f:
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.write(f)
        temp_path = f.name

    try:
        # Should not raise ValueError for unsupported format
        result = tomd(temp_path)
        assert isinstance(result, str)
    finally:
        Path(temp_path).unlink()


def test_docx_extension_supported():
    """Test that .docx extension is recognized."""
    doc = Document()
    doc.add_paragraph("Test")

    with tempfile.NamedTemporaryFile(mode="wb", suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    try:
        # Should not raise ValueError for unsupported format
        result = tomd(temp_path)
        assert isinstance(result, str)
        assert "Test" in result
    finally:
        Path(temp_path).unlink()
