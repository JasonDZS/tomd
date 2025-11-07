"""DOCX to Markdown converter."""

from pathlib import Path

from docx import Document


class DOCXConverter:
    """Convert DOCX files to Markdown."""

    def convert(self, file_path: Path) -> str:
        """
        Convert a DOCX file to Markdown.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Markdown string representation
        """
        doc = Document(file_path)
        markdown_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Handle heading styles
            if paragraph.style.name.startswith("Heading"):
                level = paragraph.style.name.replace("Heading ", "")
                try:
                    level_num = int(level)
                    markdown_parts.append(f"{'#' * level_num} {text}")
                except ValueError:
                    markdown_parts.append(text)
            else:
                # Handle text formatting
                formatted_text = self._format_runs(paragraph)
                markdown_parts.append(formatted_text)

        # Process tables
        for table in doc.tables:
            markdown_parts.append(self._convert_table(table))

        return "\n\n".join(markdown_parts).strip()

    def _format_runs(self, paragraph) -> str:
        """Format text runs with bold and italic."""
        result = []
        for run in paragraph.runs:
            text = run.text
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            result.append(text)
        return "".join(result)

    def _convert_table(self, table) -> str:
        """Convert a DOCX table to Markdown table format."""
        if not table.rows:
            return ""

        lines = []

        # Header row
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        # Data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)
